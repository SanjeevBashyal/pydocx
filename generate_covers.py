import pandas as pd
from docx import Document
from docx2pdf import convert
from pypdf import PdfWriter
import os
import sys

# Configuration
COVER_DATA_FILE = 'cover_data.xlsx'
TEMPLATE_MAIN = 'Annex cover main.docx'
TEMPLATE_SUB = 'Annex cover sub main.docx'
TEMPLATE_DESC = 'Annex cover description.docx'

def replace_text_in_paragraph(paragraph, key, value):
    """
    Replace occurrences of 'key' with 'value' in a paragraph.
    First tries to find the key in individual runs to preserve formatting.
    Falls back to text replacement if split across runs (may lose formatting).
    """
    if key in paragraph.text:
        replaced_in_run = False
        for run in paragraph.runs:
            if key in run.text:
                run.text = run.text.replace(key, value)
                replaced_in_run = True
        
        # If the key was found in paragraph.text but not in any single run,
        # it means the key is split across runs (e.g. bold/unbold boundaries).
        # In this case, we have to replace the whole text, which loses formatting.
        if not replaced_in_run:
            # Alternative: smarter run merging, but for now simple fallback.
            paragraph.text = paragraph.text.replace(key, value)

def process_template(row, template_path, output_docx_path):
    doc = Document(template_path)
    
    # Iterate over columns in the row
    for key, value in row.items():
        search_text = str(key)
        
        # Handle NaN/None and formatting
        if pd.isna(value):
            replace_text = ""
        else:
            # Check if it's a number that is effectively an integer (e.g. 1.0)
            if isinstance(value, (int, float)) and value == int(value):
                replace_text = str(int(value))
            else:
                replace_text = str(value)
                
        # Check paragraphs in body
        for p in doc.paragraphs:
            replace_text_in_paragraph(p, search_text, replace_text)
            
        # Check tables
        for table in doc.tables:
            for row_obj in table.rows:
                for cell in row_obj.cells:
                    for p in cell.paragraphs:
                        replace_text_in_paragraph(p, search_text, replace_text)
    
    doc.save(output_docx_path)

def main():
    if not os.path.exists(COVER_DATA_FILE):
        print(f"Error: {COVER_DATA_FILE} not found.")
        return
    # Templates checked later or assume existence for now

    # 1. Read Data
    try:
        # Read Excel
        df = pd.read_excel(COVER_DATA_FILE)
    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return

    print("Headers detected:", df.columns.tolist())
    
    # Check if <FILE-NAME> column exists
    file_name_col = '<FILE-NAME>'
    if file_name_col not in df.columns:
        # Try to find a column that looks like it (case insensitive or stripped)
        match = next((c for c in df.columns if c.strip() == '<FILE-NAME>'), None)
        if match:
            file_name_col = match
        else:
            print(f"Error: Column '{file_name_col}' not found in CSV.")
            return

    # Create PDFs directory
    output_dir = 'PDFs'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"Created output directory: {output_dir}")

    # 2. Generate DOCX and Convert to PDF
    print("Processing rows...")
    for index, row in df.iterrows():
        # Get filename
        file_name = row.get(file_name_col)
        if pd.isna(file_name) or str(file_name).strip() == "":
            print(f"  Warning: Row {index + 1} has empty {file_name_col}. Skipping.")
            continue
        
        file_name = str(file_name).strip()
        if not file_name.lower().endswith('.pdf'):
            file_name += '.pdf'
            
        target_pdf_path = os.path.join(output_dir, file_name)
        target_pdf_path = os.path.abspath(target_pdf_path)
        
        # Determine templates
        sub_comp_val = row.get('<SUB-COMPONENT>')
        if pd.isna(sub_comp_val) or str(sub_comp_val).strip() == "":
            # Use Main + Description
            templates_to_process = [TEMPLATE_MAIN, TEMPLATE_DESC]
        else:
            # Use Sub Main only
            templates_to_process = [TEMPLATE_SUB]
            
        generated_part_pdfs = []
        
        try:
            print(f"  Processing row {index + 1}: {file_name}")
            
            for i, tmpl in enumerate(templates_to_process):
                if not os.path.exists(tmpl):
                    print(f"    Error: Template '{tmpl}' not found. Skipping part.")
                    continue
                
                print(f"    Using template: {tmpl}")
                temp_docx = f"temp_cover_{index}_{i}.docx"
                temp_pdf = f"temp_cover_{index}_{i}.pdf"
                
                # Process docx
                process_template(row, tmpl, temp_docx)
                
                # Convert to PDF
                abs_docx = os.path.abspath(temp_docx)
                abs_pdf = os.path.abspath(temp_pdf)
                # print(f"    Converting part {i+1}...")
                convert(abs_docx, abs_pdf)
                
                if os.path.exists(abs_pdf):
                    generated_part_pdfs.append(abs_pdf)
                
                # Cleanup single docx immediately
                if os.path.exists(temp_docx):
                    try:
                        os.remove(temp_docx)
                    except:
                        pass

            # Merge if multiple parts, or rename if single
            if generated_part_pdfs:
                if len(generated_part_pdfs) == 1:
                    # Move/Rename
                    # If target exists and is same file? No, temp has different name.
                    # Remove target if exists (shutil.move might fail on windows if exists?)
                    if os.path.exists(target_pdf_path):
                        os.remove(target_pdf_path)
                    os.rename(generated_part_pdfs[0], target_pdf_path)
                else:
                    # Merge
                    print(f"    Merging {len(generated_part_pdfs)} parts to {file_name}...")
                    merger = PdfWriter()
                    for pdf in generated_part_pdfs:
                        merger.append(pdf)
                    merger.write(target_pdf_path)
                    merger.close()
                
                print(f"    Created: {target_pdf_path}")
            else:
                print("    Warning: No PDF parts generated for this row.")

        except Exception as e:
            print(f"  Error processing row {index + 1}: {e}")
        finally:
            # Cleanup temp PDFs
            for pdf in generated_part_pdfs:
                if os.path.exists(pdf):
                    try:
                        os.remove(pdf)
                    except:
                        pass

    print("Done.")

if __name__ == "__main__":
    main()
